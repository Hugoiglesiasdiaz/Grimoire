"""
Lógica de subida y extracción de texto de archivos
"""
import os
import mimetypes
from werkzeug.utils import secure_filename
from app import db
from app.models import Document
from pypdf import PdfReader
import docx


class FileService:
    """Service para manejar operaciones de archivo"""
    
    ALLOWED_EXTENSIONS = {'pdf', 'txt', 'docx', 'doc'}
    UPLOAD_FOLDER = 'uploads'
    
    def __init__(self):
        """Inicializar servicio"""
        if not os.path.exists(self.UPLOAD_FOLDER):
            os.makedirs(self.UPLOAD_FOLDER)
    
    def save_file(self, file):
        """Guardar archivo subido a disco"""
        try:
            if not self._allowed_file(file.filename):
                return {
                    'success': False,
                    'error': f'File type not allowed. Allowed types: {self.ALLOWED_EXTENSIONS}'
                }
            
            filename = secure_filename(file.filename)
            filepath = os.path.join(self.UPLOAD_FOLDER, filename)
            
            # Handle duplicate filenames
            counter = 1
            name, ext = os.path.splitext(filename)
            while os.path.exists(filepath):
                filename = f"{name}_{counter}{ext}"
                filepath = os.path.join(self.UPLOAD_FOLDER, filename)
                counter += 1
            
            file.save(filepath)
            file_size = os.path.getsize(filepath)
            
            # Create document record in database
            document = Document(
                filename=filename,
                original_filename=file.filename,
                file_path=filepath,
                file_type=self._get_file_type(filename),
                file_size=file_size
            )
            db.session.add(document)
            db.session.commit()
            
            return {
                'success': True,
                'message': 'File uploaded successfully',
                'document_id': document.id,
                'filename': filename,
                'file_size': file_size
            }
        except Exception as e:
            db.session.rollback()
            return {'success': False, 'error': str(e)}
    
    def extract_text(self, document_id):
        """Extraer texto de documento"""
        try:
            document = Document.query.get(document_id)
            if not document:
                return {'success': False, 'error': 'Document not found'}
            
            filepath = document.file_path
            text = self._extract_text_from_file(filepath, document.file_type)
            
            # Save extracted text to database
            document.extracted_text = text
            db.session.commit()
            
            return {
                'success': True,
                'message': 'Text extracted successfully',
                'text': text,
                'document_id': document_id
            }
        except Exception as e:
            db.session.rollback()
            return {'success': False, 'error': str(e)}
    
    def _allowed_file(self, filename):
        """Verificar si el archivo es permitido"""
        return '.' in filename and filename.rsplit('.', 1)[1].lower() in self.ALLOWED_EXTENSIONS
    
    def _get_file_type(self, filename):
        """Obtener tipo de archivo"""
        ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
        return ext or 'unknown'
    
    def _extract_text_from_file(self, filepath, file_type):
        """Extraer texto según el tipo de archivo"""
        if file_type == 'pdf':
            return self._extract_text_from_pdf(filepath)
        elif file_type in ['docx', 'doc']:
            return self._extract_text_from_docx(filepath)
        elif file_type == 'txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            raise ValueError(f'Unsupported file type: {file_type}')
    
    def _extract_text_from_pdf(self, filepath):
        """Extraer texto de PDF"""
        text = []
        with open(filepath, 'rb') as file:
            reader = PdfReader(file)
            for page in reader.pages:
                text.append(page.extract_text())
        return '\n'.join(text)
    
    def _extract_text_from_docx(self, filepath):
        """Extraer texto de DOCX"""
        doc = docx.Document(filepath)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return '\n'.join(text)
